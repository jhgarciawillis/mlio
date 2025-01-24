import pandas as pd
import numpy as np
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches

from core import config
from core.exceptions import MetricError
from core.state_manager import state_manager
from utils import logger
from utils.decorators import monitor_performance, handle_exceptions, log_execution
from visualization import plotter

class MetricsReporter:
    """Handle metrics reporting and visualization."""
    
    def __init__(self):
        self.report_history: List[Dict[str, Any]] = []
        self.report_templates: Dict[str, Dict[str, Any]] = {}
        self.current_report: Optional[Dict[str, Any]] = None
        
    @monitor_performance
    @handle_exceptions(MetricError)
    def generate_report(
        self,
        metrics_data: Dict[str, Any],
        report_config: Optional[Dict[str, Any]] = None,
        cluster_info: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Generate comprehensive metrics report."""
        try:
            # Record operation start
            state_monitor.record_operation_start(
                'metrics_report_generation',
                'reporting'
            )
            
            report_id = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            if report_config is None:
                report_config = self._get_default_report_config()
            
            report = {
                'metadata': {
                    'report_id': report_id,
                    'timestamp': datetime.now().isoformat(),
                    'config': report_config
                },
                'sections': {}
            }

            # Generate each section based on config
            if report_config.get('summary', True):
                report['sections']['summary'] = self._generate_summary_section(
                    metrics_data,
                    cluster_info
                )
            
            if report_config.get('detailed_metrics', True):
                report['sections']['detailed_metrics'] = self._generate_metrics_section(
                    metrics_data,
                    cluster_info
                )
            
            if report_config.get('analysis', True):
                report['sections']['analysis'] = self._generate_analysis_section(
                    metrics_data
                )
            
            if report_config.get('visualizations', True):
                report['sections']['visualizations'] = self._generate_visualization_section(
                    metrics_data,
                    cluster_info
                )
            
            if report_config.get('recommendations', True):
                report['sections']['recommendations'] = self._generate_recommendations(
                    metrics_data
                )

            # Store current report
            self.current_report = report
            
            # Record report generation
            self._record_report_generation(report_id, report_config)
            
            # Record operation completion
            state_monitor.record_operation_end(
                'metrics_report_generation',
                'completed',
                {'report_id': report_id}
            )
            
            return report
            
        except Exception as e:
            state_monitor.record_operation_end(
                'metrics_report_generation',
                'failed',
                {'error': str(e)}
            )
            raise MetricError(f"Error generating metrics report: {str(e)}") from e
    
    def _generate_summary_section(
        self,
        metrics_data: Dict[str, Any],
        cluster_info: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Generate summary section of the report."""
        summary = {
            'overall_performance': metrics_data.get('overall_metrics', {}),
            'key_findings': self._extract_key_findings(metrics_data),
            'data_overview': {
                'timestamp': datetime.now().isoformat(),
                'metrics_calculated': list(metrics_data.keys())
            }
        }
        
        if cluster_info:
            summary['cluster_overview'] = {
                'n_clusters': len(cluster_info.get('labels', [])),
                'cluster_sizes': cluster_info.get('sizes', {}),
                'cluster_performance': self._summarize_cluster_performance(
                    metrics_data.get('cluster_metrics', {})
                )
            }
        
        return summary
    
    def _generate_metrics_section(
        self,
        metrics_data: Dict[str, Any],
        cluster_info: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Generate detailed metrics section."""
        metrics_section = {
            'performance_metrics': metrics_data.get('overall_metrics', {}),
            'error_metrics': metrics_data.get('error_analysis', {}),
            'distribution_metrics': metrics_data.get('distribution_analysis', {})
        }
        
        if cluster_info:
            metrics_section['cluster_metrics'] = {
                'per_cluster': metrics_data.get('cluster_metrics', {}),
                'cluster_comparison': self._compare_cluster_performance(
                    metrics_data.get('cluster_metrics', {})
                )
            }
        
        return metrics_section
    
    def _generate_analysis_section(
        self,
        metrics_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate analysis section."""
        return {
            'error_analysis': self._analyze_error_patterns(
                metrics_data.get('error_analysis', {})
            ),
            'distribution_analysis': self._analyze_distributions(
                metrics_data.get('distribution_analysis', {})
            ),
            'performance_insights': self._generate_performance_insights(metrics_data)
        }
    
    def _generate_visualization_section(
        self,
        metrics_data: Dict[str, Any],
        cluster_info: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Generate visualization section."""
        visualizations = {}
        
        # Performance visualizations
        visualizations['performance'] = plotter.create_plot(
            'bar',
            pd.DataFrame({
                'Metric': list(metrics_data['overall_metrics'].keys()),
                'Value': list(metrics_data['overall_metrics'].values())
            }),
            x='Metric',
            y='Value',
            title='Performance Metrics'
        )
        
        # Error distribution
        if 'error_analysis' in metrics_data:
            visualizations['error_dist'] = plotter.create_plot(
                'histogram',
                pd.DataFrame({
                    'Error': metrics_data['error_analysis'].get('errors', [])
                }),
                x='Error',
                title='Error Distribution'
            )
        
        # Cluster-specific visualizations
        if cluster_info:
            cluster_metrics = metrics_data.get('cluster_metrics', {})
            if cluster_metrics:
                visualizations['cluster_performance'] = self._create_cluster_performance_plot(
                    cluster_metrics
                )
        
        return visualizations
    
    def _generate_recommendations(
        self,
        metrics_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate recommendations based on metrics."""
        recommendations = {
            'performance_improvements': [],
            'model_adjustments': [],
            'data_quality': []
        }
        
        # Analyze overall performance
        overall_metrics = metrics_data.get('overall_metrics', {})
        if overall_metrics.get('r2', 0) < 0.7:
            recommendations['model_adjustments'].append(
                "Consider using more complex model architectures"
            )
        
        # Analyze errors
        error_analysis = metrics_data.get('error_analysis', {})
        if error_analysis.get('error_statistics', {}).get('mean_error', 0) > 0:
            recommendations['model_adjustments'].append(
                "Model shows systematic bias - consider retraining with balanced data"
            )
        
        # Analyze distributions
        dist_analysis = metrics_data.get('distribution_analysis', {})
        if not dist_analysis.get('distribution_comparison', {}).get('distributions_similar', True):
            recommendations['data_quality'].append(
                "Significant difference between predicted and actual distributions"
            )
        
        return recommendations
    
    @monitor_performance
    def export_report(
        self,
        format: str = 'docx',
        output_path: Optional[Path] = None
    ) -> Path:
        """Export report in specified format."""
        if self.current_report is None:
            raise MetricError("No report available for export")
        
        if output_path is None:
            output_path = config.directories.metrics_reports
            output_path.mkdir(parents=True, exist_ok=True)
        
        export_funcs = {
            'docx': self._export_to_docx,
            'html': self._export_to_html,
            'pdf': self._export_to_pdf,
            'json': self._export_to_json
        }
        
        if format not in export_funcs:
            raise MetricError(f"Unsupported export format: {format}")
        
        return export_funcs[format](self.current_report, output_path)
    
    def _export_to_docx(
        self,
        report: Dict[str, Any],
        output_path: Path
    ) -> Path:
        """Export report to Word document."""
        doc = Document()
        
        # Add title
        doc.add_heading('Metrics Analysis Report', 0)
        
        # Add sections
        for section_name, section_content in report['sections'].items():
            doc.add_heading(section_name.replace('_', ' ').title(), level=1)
            self._add_section_to_docx(doc, section_content)
        
        # Save document
        file_path = output_path / f"metrics_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(file_path))
        
        return file_path
    
    def _get_default_report_config(self) -> Dict[str, Any]:
        """Get default report configuration."""
        return {
            'summary': True,
            'detailed_metrics': True,
            'analysis': True,
            'visualizations': True,
            'recommendations': True
        }
    
    def _record_report_generation(
        self,
        report_id: str,
        config: Dict[str, Any]
    ) -> None:
        """Record report generation in history."""
        record = {
            'timestamp': datetime.now().isoformat(),
            'report_id': report_id,
            'configuration': config
        }
        
        self.report_history.append(record)
        state_manager.set_state(
            f'metrics.reports.history.{len(self.report_history)}',
            record
        )

# Create global metrics reporter instance
metrics_reporter = MetricsReporter()